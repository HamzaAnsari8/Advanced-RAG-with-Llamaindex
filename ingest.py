print("📥 Starting ingestion...")
import os
import re
from tqdm import tqdm
from llama_index.core import Document
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.core import StorageContext, VectorStoreIndex
import chromadb
from pypdf import PdfReader
import docx

# CONFIG
DATA_DIR = "data"
VECTOR_DB_DIR = "vectordb"
TRACK_FILE = "ingested_files.txt"

EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
SUPPORTED_EXTENSIONS = [".pdf", ".txt", ".docx"]

def clean_text(text):     # clean text
    if not isinstance(text, str):
        return ""

    text = text.encode("utf-8", "ignore").decode("utf-8")
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def load_txt(filepath):      # file loader
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def load_pdf(filepath):
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def load_docx(filepath):
    document = docx.Document(filepath)
    return "\n".join([para.text for para in document.paragraphs])

def load_ingested_files():
    return set()


def save_ingested_file(filename):
    with open(TRACK_FILE, "a") as f:
        f.write(filename + "\n")

# VECTOR DB
def get_vector_store():
    # ensure vectordb folder exists
    os.makedirs(VECTOR_DB_DIR, exist_ok=True)
    chroma_client = chromadb.PersistentClient(path=VECTOR_DB_DIR)
    collection = chroma_client.get_or_create_collection(
        name="rag_collection"
    )
    vector_store = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(
        vector_store=vector_store
    )
    embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL)
    return storage_context, embed_model

# SINGLE INGETS
def ingest_single_file(filepath, filename):
    print(f"\nIngesting uploaded file: {filename}")
    ingested_files = load_ingested_files()
    if filename in ingested_files:
        print("⏭ Already ingested, skipping...")
        return

    try:
        text = ""
        if filename.lower().endswith(".txt"):
            text = load_txt(filepath)
        elif filename.lower().endswith(".pdf"):
            text = load_pdf(filepath)
        elif filename.lower().endswith(".docx"):
            text = load_docx(filepath)
        text = clean_text(text)
        if not text:
            print("Empty file")
            return
        document = Document(text=text, metadata={"file": filename})
        parser = SimpleNodeParser.from_defaults(
            chunk_size=500,
            chunk_overlap=50
        )
        nodes = parser.get_nodes_from_documents([document])
        storage_context, embed_model = get_vector_store()
        index = VectorStoreIndex(
            [],
            storage_context=storage_context,
            embed_model=embed_model
        )
        index.insert_nodes(nodes)
        save_ingested_file(filename)
        print(f"Successfully ingested: {filename}")
    except Exception as e:
        print(f"Error ingesting {filename}: {e}")

# BULK LOAD
def load_documents():
    documents = []
    ingested_files = load_ingested_files()
    new_files_count = 0

    if not os.path.exists(DATA_DIR):
        print("Data folder not found.")
        return documents

    files = os.listdir(DATA_DIR)
    for filename in tqdm(files, desc="Processing data folder"):
        filepath = os.path.join(DATA_DIR, filename)

        try:
            if not any(filename.lower().endswith(ext) for ext in SUPPORTED_EXTENSIONS):
                continue

            if filename in ingested_files:
                continue
            text = ""
            if filename.lower().endswith(".txt"):
                text = load_txt(filepath)
            elif filename.lower().endswith(".pdf"):
                text = load_pdf(filepath)
            elif filename.lower().endswith(".docx"):
                text = load_docx(filepath)
            text = clean_text(text)
        
            if text:
                documents.append(
                    Document(text=text, metadata={"file": filename})
                )
                save_ingested_file(filename)
                new_files_count += 1

        except Exception as e:
            print(f"Error: {filename} -> {e}")

    print(f"\nNew files added: {new_files_count}")
    return documents

# BULK INGEST
def ingest():
    documents = load_documents()
    if not documents:
        print("No new documents to ingest!")
        return
    parser = SimpleNodeParser.from_defaults(
        chunk_size=500,
        chunk_overlap=50
    )
    nodes = parser.get_nodes_from_documents(documents)
    storage_context, embed_model = get_vector_store()
    index = VectorStoreIndex(
        [],
        storage_context=storage_context,
        embed_model=embed_model
    )
    index.insert_nodes(nodes)
    print("Ingestion completed!")

# RUN
if __name__ == "__main__":
    ingest()